from docx import Document
from docx.shared import Pt , Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT , WD_BREAK

from random import shuffle

import pandas as pd
import pyqrcode


def make_quizz_doc(nbq , nbt , quizzpath):

    """
    makes the .doc    
    """
    
    df = pd.read_excel(quizzpath + "/quizz.xlsx")  
    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    questions = []
    for i in range(2,nbq+2):
        question = i-1 , df["questions"][i-2], [df[k][i-2] for k in ["prop1","prop2","prop3","prop4"]] 
        questions.append(question)

     
    for i in range(1,nbt+1):
        code = ''
        shuffle(questions)
        #make qrcode for shuffled sequence
        code = ' '.join(str(num) for num, _ ,_ in questions)
        url = pyqrcode.create(code)
        url.png('code.png' , scale=6 , module_color=[0, 0, 0, 128], background=[0xff, 0xff, 0xff])        

        #header with a table 1 row 2 columns
        table = document.add_table(rows=1, cols=2)
        #left cell
        cell = table.cell(0,0)
        cell.text = 'Nom: \nPrénom:\nClasse:'
        #right cell
        cell = table.cell(0,1)
        p = cell.paragraphs[0]
        p.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture('code.png',width=Cm(3.75))

        #generate text for multiple choice questions 
        table = document.add_table(rows=nbq, cols=2)
       
        for i , question in enumerate(questions):
            num , qtext , props = question
            table.cell(i,0).text = qtext 
            cell = table.cell(i,0)
            test= [str(i+1)+")"+str(props[i]) for i in range(0,len(props))]
            
            ptext = '\n'.join(test)                    
            table.cell(i,1).text = ptext 
            
        
        #go to next page
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  

        #make bubble sheets


    document.save(quizzpath+'/quizztest.docx')