from docx import Document , section
from docx.shared import Pt , Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT , WD_BREAK 
from docx.enum.section import WD_SECTION
from fpdf import FPDF

from random import shuffle

import openpyxl
import pyqrcode


def make_quizz_doc(nbq , nbt , quizzpath):
    """makes the .doc    
    """
    book = openpyxl.load_workbook(quizzpath+'/quizz.xlsx')
    sheet = book.active    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    questions = []
    for i in range(2,nbq+2):
        question = i-1 , sheet.cell(row=i,column=1).value , [sheet.cell(row=i,column=j).value for j in range(2,6) ] 
        questions.append(question)
     
    for i in range(1,nbt+1):
        code = ''
        shuffle(questions)
        #make qrcode for shuffled sequence
        code = ' '.join(str(num) for num, _ ,_ in questions)
        url = pyqrcode.create(code)
        url.png('code.png' , scale=6 , module_color=[0, 0, 0, 128], background=[0xff, 0xff, 0xff])        

        #header with a table 1 row 2 columns
        table = document.add_table(rows=1, cols=2)
        #left cell
        cell = table.cell(0,0)
        cell.text = 'Nom: \nPrénom:\nClasse:'
        #right cell
        cell = table.cell(0,1)
        p = cell.paragraphs[0]
        p.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture('code.png',width=Cm(3.75))

        #generate text for multiple choice questions 
        table = document.add_table(rows=nbq, cols=2)
       
        for i , question in enumerate(questions):
            num , qtext , props = question
            table.cell(i,0).text = qtext 
            cell = table.cell(i,0)
            test= [str(i+1)+")"+props[i] for i in range(0,len(props))]
            print(test)
            ptext = '\n'.join(test)                    
            table.cell(i,1).text = ptext 
            
        
        #go to next page
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)  

        #make bubble sheets


    document.save(quizzpath+'/quizztest.docx')

def make_bubble_sheet(nb,first,second=None):
    document = Document("little sheet.docx")
    qrcodesection = document.sections[-1]
    
    #table = document.add_table(rows=1, cols=2)
    #cell = table.cell(0,0)
    #p = cell.paragraphs[0]
    #p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    #p.add_run().add_picture('code.png',width=Cm(4.5))
    p = document.add_paragraph()
    p.alignment=WD_PARAGRAPH_ALIGNMENT.CENTER
    p.add_run().add_picture(first,width=Cm(4.5))
    p.add_run().add_text("\t\t\t\t\t\t")
    if second != None:
        p.add_run().add_picture(second,width=Cm(4.5))
    document.save("page bubble sheet.docx")
       
    

#make_quizz_doc(10 ,2 ,"E:/work/python/quizz scanner" )
make_bubble_sheet(30)