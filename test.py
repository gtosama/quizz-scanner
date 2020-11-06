from docx import Document
from docx.shared import Pt , Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT , WD_BREAK
from fpdf import FPDF

from random import shuffle

import openpyxl
import pyqrcode


def make_quizz_doc(nbq , nbt , quizzpath):
    """makes the .doc    
    """
    book = openpyxl.load_workbook(quizzpath+'/quizz.xlsx')
    sheet = book.active    
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(12)
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin = Cm(1)
        section.right_margin = Cm(1)

    questions = []
    for i in range(2,nbq+2):
        question = i-1 , sheet.cell(row=i,column=1).value , [sheet.cell(row=i,column=j).value for j in range(2,6) ] 
        questions.append(question)
     
    for i in range(1,nbt+1):
        code = ''
        shuffle(questions)
        #make qrcode for shuffled sequence
        code = ' '.join(str(num) for num, _ ,_ in questions)
        url = pyqrcode.create(code)
        url.png('code.png' , scale=6 , module_color=[0, 0, 0, 128], background=[0xff, 0xff, 0xff])        

        #header with a table 1 row 2 columns
        table = document.add_table(rows=1, cols=2)
        #left cell
        cell = table.cell(0,0)
        cell.text = 'Nom: \nPrénom:\nClasse:'
        #right cell
        cell = table.cell(0,1)
        p = cell.paragraphs[0]
        p.alignment=WD_PARAGRAPH_ALIGNMENT.RIGHT
        p.add_run().add_picture('code.png',width=Cm(3.75))

        #generate text for multiple choice questions 
        table = document.add_table(rows=10, cols=2)
       
        for i , question in enumerate(questions):
            num , qtext , props = question 
            table.cell(i,0).text = qtext               
            for pr in props:              
                if pr != None:
                    table.cell(i,1).add_paragraph( pr , style='List Bullet')       
        
        #go to next page
        document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)    


    document.save(quizzpath+'/quizztest.docx')

def make_bubble_sheet(nbq):
    document = Document()
    table = document.add_table(rows=1 , cols=1)
    table.style = document.styles['Table Grid']
    

make_quizz_doc(10 ,5 ,"E:/work/python/quizz scanner" )